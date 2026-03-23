import hashlib
import logging
import re
from pathlib import Path
from typing import Dict, Iterable, List, Tuple
from xml.etree import ElementTree
from zipfile import ZipFile

import pandas as pd
from docx import Document
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from .db import Base, SessionLocal, engine, ensure_database_schema
from .llm import chat_completion, llm_enabled
from .models import Chunk, Doc
from .utils_text import chunk_text, normalize_whitespace

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".dotx", ".xlsx", ".xls", ".csv"}


def _file_checksum(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file_handle:
        for block in iter(lambda: file_handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _pdf_to_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _document_to_text(path: Path) -> str:
    try:
        document = Document(str(path))
        parts: List[str] = []

        for paragraph in document.paragraphs:
            line = paragraph.text.strip()
            if line:
                parts.append(line)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except ValueError:
        return _openxml_template_to_text(path)


def _openxml_template_to_text(path: Path) -> str:
    with ZipFile(path) as archive:
        xml_bytes = archive.read("word/document.xml")

    root = ElementTree.fromstring(xml_bytes)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: List[str] = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [
            node.text.strip()
            for node in paragraph.findall(".//w:t", namespace)
            if node.text and node.text.strip()
        ]
        if texts:
            lines.append(" ".join(texts))
    return "\n".join(lines)


def _normalize_column_name(value: str, fallback_index: int) -> str:
    candidate = str(value or "").strip()
    if not candidate or candidate.lower().startswith("unnamed"):
        return f"column_{fallback_index + 1}"
    return re.sub(r"\s+", " ", candidate)


def _normalize_cell(value: str) -> str:
    return re.sub(r"\s+", " ", str(value).strip())


def _dataframe_to_text(frame: pd.DataFrame, sheet_name: str) -> str:
    if frame.empty:
        return ""

    working = frame.copy()
    working = working.dropna(axis=0, how="all").dropna(axis=1, how="all")
    if working.empty:
        return ""

    working.columns = [
        _normalize_column_name(column, index) for index, column in enumerate(working.columns)
    ]
    working = working.fillna("").astype(str)

    parts: List[str] = [f"Лист {sheet_name}"]
    named_columns = [column for column in working.columns if not column.startswith("column_")]

    for row_number, (_, row) in enumerate(working.iterrows(), start=1):
        values = [_normalize_cell(value) for value in row.tolist()]
        if not any(values):
            continue

        row_parts: List[str] = []
        for column, value in zip(working.columns, values):
            if not value:
                continue
            if named_columns:
                row_parts.append(f"{column}: {value}")
            else:
                row_parts.append(value)

        if row_parts:
            parts.append(f"Строка {row_number}: " + "; ".join(row_parts))

    return "\n".join(parts)


def _excel_to_text(path: Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts: List[str] = []
    for sheet_name, frame in sheets.items():
        sheet_text = _dataframe_to_text(frame, str(sheet_name))
        if sheet_text:
            parts.append(sheet_text)
    return "\n\n".join(parts)


def _csv_to_text(path: Path) -> str:
    last_error = None
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            frame = pd.read_csv(path, dtype=str, encoding=encoding)
            return _dataframe_to_text(frame, path.stem)
        except UnicodeDecodeError as exc:
            last_error = exc
    if last_error:
        raise last_error
    return ""


def _extract_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _pdf_to_text(path)
    if extension in {".docx", ".dotx"}:
        return _document_to_text(path)
    if extension in {".xlsx", ".xls"}:
        return _excel_to_text(path)
    if extension == ".csv":
        return _csv_to_text(path)
    raise ValueError(f"Unsupported file type: {extension}")


def _iter_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path.resolve()


def _heuristic_description(title: str, file_type: str, content: str) -> str:
    preview = normalize_whitespace(content)[:500]
    first_sentence = preview.split(".")[0].strip()
    if file_type in {"xls", "xlsx", "csv"}:
        return (
            f"Что внутри: табличный документ '{title}' с расписанием, перечнем строк или структурированными данными. "
            f"Когда обращаться: когда нужно найти даты, занятия, экзамены, группы или конкретные позиции в таблице. "
            f"Краткая выжимка: {first_sentence}"
        )
    if file_type in {"pdf", "docx", "dotx"}:
        return (
            f"Что внутри: документ '{title}' с правилами, регламентом или шаблоном. "
            f"Когда обращаться: когда нужно понять требования, порядок действий или официальный формат. "
            f"Краткая выжимка: {first_sentence}"
        )
    return (
        f"Что внутри: документ '{title}'. Когда обращаться: когда нужен официальный контекст по теме документа. "
        f"Краткая выжимка: {first_sentence}"
    )


def _generate_description(title: str, file_type: str, content: str) -> str:
    excerpt = normalize_whitespace(content)[:5000]
    if not llm_enabled():
        return _heuristic_description(title, file_type, excerpt)

    prompt = (
        f"Название документа: {title}\n"
        f"Тип файла: {file_type}\n"
        "Нужно написать краткое описание документа на русском языке в 2-4 предложениях.\n"
        "Формат:\n"
        "Что внутри: ...\n"
        "Когда обращаться: ...\n\n"
        f"Фрагмент документа:\n{excerpt}"
    )
    description = chat_completion(
        "Ты анализируешь документы учебного офиса и кратко объясняешь их назначение.",
        prompt,
        temperature=0.1,
        max_tokens=220,
    )
    if description:
        return description
    return _heuristic_description(title, file_type, excerpt)


def _upsert_doc(
    session: Session,
    file_path: Path,
    checksum: str,
) -> Tuple[Doc, str]:
    resolved_path = str(file_path.resolve())
    existing = session.execute(
        select(Doc).where(Doc.file_path == resolved_path)
    ).scalar_one_or_none()
    title = file_path.stem
    file_type = file_path.suffix.lower().lstrip(".")

    if existing:
        existing.title = title
        existing.file_type = file_type
        existing.status = "active"

        if existing.checksum == checksum:
            if existing.description.strip():
                return existing, "skipped"
            return existing, "metadata"

        existing.checksum = checksum
        return existing, "updated"

    document = Doc(
        title=title,
        description="",
        file_path=resolved_path,
        file_type=file_type,
        checksum=checksum,
        status="active",
    )
    session.add(document)
    session.flush()
    return document, "new"


def index_path(path: str) -> Dict[str, int]:
    Base.metadata.create_all(bind=engine)
    ensure_database_schema()
    counts = {"indexed_docs": 0, "updated_docs": 0, "skipped_docs": 0}

    root = Path(path).resolve()
    if not root.exists() or not root.is_dir():
        logger.warning("Data directory not found: %s", root)
        return counts

    with SessionLocal() as session:
        for file_path in _iter_files(root):
            try:
                checksum = _file_checksum(file_path)
                doc, action = _upsert_doc(session, file_path, checksum)
                if action == "skipped":
                    counts["skipped_docs"] += 1
                    session.commit()
                    continue

                content = normalize_whitespace(_extract_text(file_path))
                if not content:
                    logger.warning("No extractable text in %s", file_path)
                    session.rollback()
                    counts["skipped_docs"] += 1
                    continue

                doc.description = _generate_description(doc.title, doc.file_type, content)

                if action == "metadata":
                    session.commit()
                    counts["updated_docs"] += 1
                    logger.info("Refreshed metadata for %s", file_path.name)
                    continue

                if action == "updated":
                    session.query(Chunk).filter(Chunk.doc_id == doc.id).delete(
                        synchronize_session=False
                    )

                chunks = chunk_text(content)
                if not chunks:
                    logger.warning("No chunks produced for %s", file_path)
                    session.rollback()
                    counts["skipped_docs"] += 1
                    continue

                for index, chunk in enumerate(chunks):
                    session.add(
                        Chunk(
                            doc_id=doc.id,
                            chunk_index=index,
                            content=chunk,
                        )
                    )

                session.commit()
                counts["updated_docs" if action == "updated" else "indexed_docs"] += 1
                logger.info(
                    "Indexed %s as %s with %s chunks",
                    file_path.name,
                    action,
                    len(chunks),
                )
            except Exception as exc:
                session.rollback()
                counts["skipped_docs"] += 1
                logger.exception("Failed to index %s: %s", file_path, exc)

    return counts
